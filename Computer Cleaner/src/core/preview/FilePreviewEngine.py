from __future__ import annotations

import mimetypes
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff", ".tif"}
TEXT_EXTENSIONS = {".txt", ".md", ".log", ".py", ".json", ".yaml", ".yml", ".csv"}
DOC_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}
EXECUTABLE_EXTENSIONS = {".exe", ".dll", ".bat", ".cmd", ".msi", ".com", ".scr", ".ps1", ".sh"}


@dataclass
class ProcessingStats:
    processed_count: int = 0
    file_types: Counter[str] | None = None
    preview_success_count: int = 0
    extraction_success_count: int = 0
    failures: list[dict[str, str]] | None = None

    def __post_init__(self) -> None:
        if self.file_types is None:
            self.file_types = Counter()
        if self.failures is None:
            self.failures = []


class FilePreviewEngine:
    def __init__(
        self,
        *,
        thumbnail_dir: Path,
        preview_dir: Path,
        text_preview_chars: int = 5000,
        extracted_content_chars: int = 15000,
    ) -> None:
        self.thumbnail_dir = thumbnail_dir
        self.preview_dir = preview_dir
        self.text_preview_chars = text_preview_chars
        self.extracted_content_chars = extracted_content_chars

        self.thumbnail_dir.mkdir(parents=True, exist_ok=True)
        self.preview_dir.mkdir(parents=True, exist_ok=True)

    def process_file(self, file_path: Path) -> dict[str, Any]:
        metadata = self._extract_metadata(file_path)
        file_type = self._detect_file_type(file_path)

        preview: dict[str, Any] = {
            "thumbnail": None,
            "text_preview": None,
            "full_preview_available": False,
        }
        extracted_content: str | None = None

        if file_type == "image":
            preview["thumbnail"] = self._create_image_thumbnail(file_path)
            preview["full_preview_available"] = True
        elif file_type == "pdf":
            preview["thumbnail"] = self._create_pdf_thumbnail(file_path)
            extracted_content = self._extract_pdf_text(file_path)
            preview["text_preview"] = self._truncate(extracted_content, self.text_preview_chars)
            preview["full_preview_available"] = preview["thumbnail"] is not None
        elif file_type == "docx":
            extracted_content = self._extract_docx_text(file_path)
            preview["text_preview"] = self._truncate(extracted_content, self.text_preview_chars)
            preview["full_preview_available"] = bool(extracted_content)
        elif file_type == "text":
            extracted_content = self._extract_text_file(file_path)
            preview["text_preview"] = self._truncate(extracted_content, self.text_preview_chars)
            preview["full_preview_available"] = bool(extracted_content)

        if extracted_content:
            extracted_content = self._truncate(extracted_content, self.extracted_content_chars)

        return {
            "file_path": str(file_path),
            "file_type": file_type,
            "metadata": metadata,
            "preview": preview,
            "extracted_content": extracted_content,
        }

    def _detect_file_type(self, path: Path) -> str:
        mime = self._detect_mime(path)
        suffix = path.suffix.lower()

        if mime and mime.startswith("image/") or suffix in IMAGE_EXTENSIONS:
            return "image"
        if mime == "application/pdf" or suffix in PDF_EXTENSIONS:
            return "pdf"
        if suffix in DOC_EXTENSIONS or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return "docx"
        if mime and mime.startswith("text/") or suffix in TEXT_EXTENSIONS:
            return "text"
        return "unknown"

    def _detect_mime(self, path: Path) -> str | None:
        try:
            import magic  # type: ignore

            return magic.from_file(str(path), mime=True)
        except Exception:
            guessed, _ = mimetypes.guess_type(str(path))
            return guessed

    def _extract_metadata(self, path: Path) -> dict[str, Any]:
        stat = path.stat()
        created_ts = getattr(stat, "st_ctime", None)
        modified_ts = getattr(stat, "st_mtime", None)

        return {
            "file_path": str(path),
            "filename": path.name,
            "extension": path.suffix.lower().lstrip("."),
            "size_mb": round(stat.st_size / (1024 * 1024), 4),
            "created_time": self._safe_timestamp(created_ts),
            "modified_time": self._safe_timestamp(modified_ts),
            "folder_path": str(path.parent),
            "file_category_guess": self._detect_file_type(path),
        }

    @staticmethod
    def _safe_timestamp(ts: Any) -> str | None:
        if ts is None:
            return None
        try:
            return datetime.fromtimestamp(float(ts)).isoformat()
        except Exception:
            return None

    @staticmethod
    def _truncate(value: str | None, max_chars: int) -> str | None:
        if not value:
            return value
        return value[:max_chars]

    def _create_image_thumbnail(self, path: Path, size: tuple[int, int] = (512, 512)) -> str | None:
        out_path = self.thumbnail_dir / f"{path.stem}.img.thumb.png"
        try:
            from PIL import Image

            with Image.open(path) as image:
                image.thumbnail(size)
                image.convert("RGBA").save(out_path, format="PNG")
            return str(out_path)
        except Exception:
            return None

    def _create_pdf_thumbnail(self, path: Path, zoom: float = 1.5) -> str | None:
        out_path = self.thumbnail_dir / f"{path.stem}.pdf.thumb.png"
        try:
            import fitz  # type: ignore

            document = fitz.open(str(path))
            try:
                page = document.load_page(0)
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                pix.save(str(out_path))
            finally:
                document.close()
            return str(out_path)
        except Exception:
            return None

    def _extract_pdf_text(self, path: Path) -> str | None:
        try:
            import fitz  # type: ignore

            document = fitz.open(str(path))
            chunks: list[str] = []
            try:
                max_pages = min(10, len(document))
                for page_index in range(max_pages):
                    chunks.append(document.load_page(page_index).get_text("text"))
            finally:
                document.close()
            text = "\n".join(chunks).strip()
            return text or None
        except Exception:
            return None

    def _extract_docx_text(self, path: Path) -> str | None:
        try:
            import docx  # type: ignore

            document = docx.Document(str(path))
            parts: list[str] = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text.strip())

            for table in document.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append(" | ".join(row_cells))

            text = "\n".join(parts).strip()
            return text or None
        except Exception:
            return None

    def _extract_text_file(self, path: Path) -> str | None:
        try:
            return path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            return None


def is_hidden_file(path: Path) -> bool:
    if path.name.startswith("."):
        return True

    try:
        import os
        import stat as statmod

        if os.name == "nt":
            attrs = os.stat(str(path)).st_file_attributes  # type: ignore[attr-defined]
            return bool(attrs & statmod.FILE_ATTRIBUTE_HIDDEN)  # type: ignore[attr-defined]
    except Exception:
        return False

    return False


def is_system_or_executable(path: Path) -> bool:
    normalized = str(path).lower().replace("\\", "/")
    blocked_markers = [
        "/windows/",
        "/program files/",
        "/program files (x86)/",
        "/system32/",
    ]

    if any(marker in normalized for marker in blocked_markers):
        return True

    if path.suffix.lower() in EXECUTABLE_EXTENSIONS:
        return True

    return False


def build_processing_report(stats: ProcessingStats) -> dict[str, Any]:
    processed = max(stats.processed_count, 0)
    preview_success_rate = (stats.preview_success_count / processed * 100.0) if processed else 0.0
    extraction_success_rate = (stats.extraction_success_count / processed * 100.0) if processed else 0.0

    return {
        "processed_summary": {
            "number_of_files": processed,
            "file_types_breakdown": dict(stats.file_types or {}),
        },
        "success_failure": {
            "preview_success_rate_percent": round(preview_success_rate, 2),
            "extraction_success_rate_percent": round(extraction_success_rate, 2),
            "failures": stats.failures or [],
        },
        "system_evaluation": {
            "what_worked_well": [
                "Unified preview output format for all files.",
                "Safe scanning that avoids hidden/system/executable files by default.",
            ],
            "what_is_missing": [
                "OCR for scanned PDFs/images is not enabled in this module.",
                "No persistent processing queue database integration yet.",
            ],
            "bottlenecks": [
                "Large PDFs and huge text files can still be expensive to parse.",
                "Deep recursive scans are currently single-threaded.",
            ],
        },
        "improvement_suggestions": {
            "performance_improvements": [
                "Add concurrent workers for independent file processing.",
                "Cache file hashes and skip unchanged files during re-runs.",
            ],
            "better_file_handling": [
                "Add explicit support for XLSX, PPTX, and archive files.",
                "Add robust encoding detection for non-UTF8 text files.",
            ],
            "preview_enhancements": [
                "Generate multi-page PDF previews and richer document snippets.",
                "Create lower-resolution and high-resolution thumbnail variants.",
            ],
            "ai_integration_readiness": [
                "Persist extracted content + metadata to training-ready records.",
                "Attach labels/actions feedback for downstream classifier training.",
            ],
        },
    }
